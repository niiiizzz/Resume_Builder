"""
Resume export utilities — generate PDF and DOCX from markdown text.
Business logic only — no Streamlit dependency.
"""

import io
import re
import logging

logger = logging.getLogger(__name__)


def _parse_markdown_lines(text: str) -> list[dict]:
    """
    Parse markdown text into a list of structured blocks.
    Each block: {'type': 'heading'|'bullet'|'paragraph', 'level': int, 'text': str}
    """
    blocks = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append({"type": "heading", "level": level, "text": heading_match.group(2).strip()})
            continue

        # Bullet points
        bullet_match = re.match(r"^[-*•]\s+(.+)$", stripped)
        if bullet_match:
            blocks.append({"type": "bullet", "level": 0, "text": bullet_match.group(1).strip()})
            continue

        # Bold markers cleanup for plain text
        blocks.append({"type": "paragraph", "level": 0, "text": stripped})

    return blocks


def _clean_markdown(text: str) -> str:
    """Remove markdown formatting characters for plain text rendering."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # bold
    text = re.sub(r"\*(.+?)\*", r"\1", text)       # italic
    text = re.sub(r"`(.+?)`", r"\1", text)          # inline code
    return text.strip()


def resume_to_pdf(markdown_text: str) -> bytes:
    """
    Convert markdown resume text to a clean PDF.

    Args:
        markdown_text: Resume content in markdown format.

    Returns:
        PDF file as raw bytes.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    style_h1 = ParagraphStyle(
        "ResumeH1", parent=styles["Heading1"],
        fontSize=18, spaceAfter=6, textColor=HexColor("#1a1a1a"),
    )
    style_h2 = ParagraphStyle(
        "ResumeH2", parent=styles["Heading2"],
        fontSize=14, spaceAfter=4, spaceBefore=10, textColor=HexColor("#333333"),
    )
    style_h3 = ParagraphStyle(
        "ResumeH3", parent=styles["Heading3"],
        fontSize=12, spaceAfter=3, spaceBefore=6, textColor=HexColor("#444444"),
    )
    style_body = ParagraphStyle(
        "ResumeBody", parent=styles["Normal"],
        fontSize=10, leading=14, spaceAfter=3, textColor=HexColor("#222222"),
    )
    style_bullet = ParagraphStyle(
        "ResumeBullet", parent=style_body,
        leftIndent=20, bulletIndent=8, spaceBefore=1, spaceAfter=1,
    )

    heading_styles = {1: style_h1, 2: style_h2, 3: style_h3, 4: style_h3}

    elements = []
    blocks = _parse_markdown_lines(markdown_text)

    for block in blocks:
        text = _clean_markdown(block["text"])
        if block["type"] == "heading":
            style = heading_styles.get(block["level"], style_h2)
            elements.append(Paragraph(text, style))
            if block["level"] == 1:
                elements.append(HRFlowable(
                    width="100%", thickness=1,
                    color=HexColor("#00B8D9"), spaceAfter=6,
                ))
        elif block["type"] == "bullet":
            elements.append(Paragraph(f"•  {text}", style_bullet))
        else:
            elements.append(Paragraph(text, style_body))

    if not elements:
        elements.append(Paragraph("No resume content provided.", style_body))

    doc.build(elements)
    return buffer.getvalue()


def resume_to_docx(markdown_text: str) -> bytes:
    """
    Convert markdown resume text to a DOCX document.
    Preserves headings, bullet points, and paragraph structure.

    Args:
        markdown_text: Resume content in markdown format.

    Returns:
        DOCX file as raw bytes.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(34, 34, 34)

    blocks = _parse_markdown_lines(markdown_text)

    heading_level_map = {1: 0, 2: 1, 3: 2, 4: 3}

    for block in blocks:
        text = _clean_markdown(block["text"])

        if block["type"] == "heading":
            level = heading_level_map.get(block["level"], 1)
            heading = doc.add_heading(text, level=level)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(26, 26, 26)

        elif block["type"] == "bullet":
            para = doc.add_paragraph(text, style="List Bullet")
            for run in para.runs:
                run.font.size = Pt(10)

        else:
            doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
